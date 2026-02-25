"""Genera la guía de fotos de auditoría en formato Word (.docx)."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = "GUIA_FOTOS_AUDITORIA.docx"

RED = RGBColor(0xE6, 0x39, 0x46)
DARK_BLUE = RGBColor(0x1D, 0x35, 0x57)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x2E, 0xCC, 0x71)
ORANGE = RGBColor(0xF3, 0x9C, 0x12)
LIGHT_GRAY = RGBColor(0xF0, 0xF0, 0xF0)


def set_cell_shading(cell, color_hex: str):
    shading = cell._tc.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(s)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1D3557"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8F9FA")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_item_block(doc, item_id, item_name, photos, tip=None):
    p = doc.add_paragraph()
    run = p.add_run(f"{item_id}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED
    run2 = p.add_run(item_name)
    run2.bold = True
    run2.font.size = Pt(11)
    run2.font.color.rgb = DARK_BLUE
    p.space_before = Pt(12)
    p.space_after = Pt(4)

    count_p = doc.add_paragraph()
    count_run = count_p.add_run(f"Fotos a tomar ({len(photos)}):")
    count_run.bold = True
    count_run.font.size = Pt(9)
    count_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    count_p.space_after = Pt(2)

    for i, photo in enumerate(photos, 1):
        lp = doc.add_paragraph(style="List Number")
        lr = lp.add_run(photo)
        lr.font.size = Pt(9)
        lp.space_before = Pt(0)
        lp.space_after = Pt(1)
        lp.paragraph_format.left_indent = Cm(1.5)

    if tip:
        tp = doc.add_paragraph()
        tr = tp.add_run(tip)
        tr.italic = True
        tr.font.size = Pt(8)
        tr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        tp.space_before = Pt(2)


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── Portada ──
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GRIDO AUDIT VISION")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RED

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Guía de fotos para auditoría interna")
    r2.font.size = Pt(16)
    r2.font.color.rgb = DARK_BLUE

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = ref.add_run("Referencia: Auditorías Operativas Abril 2025")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    for _ in range(3):
        doc.add_paragraph()

    instr = doc.add_paragraph()
    instr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = instr.add_run(
        "Para cada ítem se listan las fotos que el colaborador debe tomar.\n"
        "Todas las fotos deben ser claras, bien iluminadas y sin filtros.\n"
        "Nombrar cada archivo con el código del ítem + número de foto.\n"
        "Ejemplo: A1_01.jpg, A1_02.jpg, etc."
    )
    ri.font.size = Pt(9)
    ri.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Plan de Periodicidad ──
    h = doc.add_heading("Plan de Periodicidad", level=1)
    for run in h.runs:
        run.font.color.rgb = RED

    doc.add_paragraph(
        "La auditoría oficial comercial es cada 6 meses. "
        "La auditoría interna se organiza en dos modalidades:"
    )

    add_styled_table(doc,
        ["Modalidad", "Frecuencia", "Alcance", "Fotos aprox."],
        [
            ["COMPLETA", "Cada 2 meses (mes 2, mes 4)", "Secciones A + B + C + D + E", "~250"],
            ["RÁPIDA", "Meses intermedios (mes 1, 3, 5)", "Sección C + desvíos anteriores", "~80-120"],
            ["OFICIAL", "Cada 6 meses (mes 6)", "Auditoría formal de la marca", "—"],
        ],
        col_widths=[3, 5, 6, 2.5],
        header_color="E63946",
    )

    doc.add_paragraph()

    h2 = doc.add_heading("Periodicidad por sección", level=2)
    for run in h2.runs:
        run.font.color.rgb = DARK_BLUE

    add_styled_table(doc,
        ["Sección", "Periodicidad", "Tipo de auditoría"],
        [
            ["A — Infraestructura", "Cada 2 meses", "Solo COMPLETA"],
            ["B — Experiencia del cliente", "Cada 2 meses", "Solo COMPLETA"],
            ["C — Operatoria diaria", "MENSUAL", "COMPLETA y RÁPIDA"],
            ["D — Imagen y formato", "Cada 2 meses", "Solo COMPLETA"],
            ["E — Oferta y stock", "Cada 2 meses", "Solo COMPLETA"],
        ],
        col_widths=[6, 4, 5],
    )

    p_why = doc.add_paragraph()
    p_why.space_before = Pt(8)
    rw = p_why.add_run(
        "¿Por qué la sección C es mensual? "
    )
    rw.bold = True
    rw.font.size = Pt(9)
    p_why.add_run(
        "Porque los ítems de operatoria diaria (higiene, manipulación, uniformes, "
        "temperaturas, almacenamiento, documentación) son los más variables y donde "
        "se concentra la mayoría de los \"No Conforme\". Revisarlos cada mes permite "
        "corregir antes de que se vuelvan hábitos."
    ).font.size = Pt(9)

    p_seg = doc.add_paragraph()
    rs = p_seg.add_run("Seguimiento de desvíos: ")
    rs.bold = True
    rs.font.size = Pt(9)
    p_seg.add_run(
        "En la auditoría rápida (meses 1, 3 y 5), además de la sección C completa, "
        "se deben re-fotografiar TODOS los ítems que resultaron \"Observación\" o "
        "\"No Conforme\" en la auditoría anterior, sin importar a qué sección pertenezcan."
    ).font.size = Pt(9)

    doc.add_paragraph()

    h3 = doc.add_heading("Calendario anual sugerido", level=2)
    for run in h3.runs:
        run.font.color.rgb = DARK_BLUE

    add_styled_table(doc,
        ["Mes", "Tipo", "Alcance", "Fotos"],
        [
            ["Ene", "RÁPIDA", "Sección C + desvíos", "~80"],
            ["Feb", "COMPLETA", "A + B + C + D + E", "~250"],
            ["Mar", "RÁPIDA", "Sección C + desvíos", "~80"],
            ["Abr", "COMPLETA", "A + B + C + D + E", "~250"],
            ["May", "RÁPIDA", "Sección C + desvíos", "~80"],
            ["Jun", "OFICIAL", "Auditoría formal de la marca", "—"],
            ["Jul", "RÁPIDA", "Sección C + desvíos", "~80"],
            ["Ago", "COMPLETA", "A + B + C + D + E", "~250"],
            ["Sep", "RÁPIDA", "Sección C + desvíos", "~80"],
            ["Oct", "COMPLETA", "A + B + C + D + E", "~250"],
            ["Nov", "RÁPIDA", "Sección C + desvíos", "~80"],
            ["Dic", "OFICIAL", "Auditoría formal de la marca", "—"],
        ],
        col_widths=[2, 3, 7, 2],
    )

    doc.add_page_break()

    # ── SECCIONES ──
    sections_data = [
        ("A", "INFRAESTRUCTURA: Estado de conservación y limpieza",
         "Cada 2 MESES (solo en auditoría COMPLETA)", [
            ("A.1", "Exterior (pisos, zócalos, cordón cuneta, techos, muros, iluminación, pérgolas, toldos, bicicletero, dog parking, juegos infantiles)", [
                "Vista general de la vereda/exterior del local (desde la calle, que se vea el frente completo incluyendo pisos y cordón cuneta).",
                "Pisos exteriores en detalle: juntas, rejillas, deck (si aplica). Acercarse para que se vean las condiciones de las piezas.",
                "Muros y zócalos exteriores: foto de cada pared exterior visible.",
                "Techos y cielorrasos exteriores (pérgola, toldo). Foto mirando hacia arriba desde abajo.",
                "Bicicletero / dog parking / juegos infantiles exteriores (si existen). Una foto de cada uno mostrando estado general.",
                "Iluminación exterior: foto de los artefactos de luz (reflectores, lámparas LED). Si es de noche, foto con luces encendidas.",
            ], "Si hay defectos visibles (graffitis, piezas rotas, manchas), sacar foto extra en detalle."),
            ("A.2", "Marquesina: Estado y limpieza", [
                "Vista frontal completa de la marquesina (que se vea toda la estructura, chapas, logo corpóreo).",
                "Detalle del logo corpóreo y su iluminación.",
                "Cartel bandera y/o saliente (si existe). Foto lateral.",
                "Detalle de limpieza: acercar a la marquesina para ver polvo, manchas, óxido o suciedad acumulada.",
            ], "Si es horario nocturno: foto adicional con reflectores encendidos."),
            ("A.3", "Mobiliario: Estado y limpieza (mesas, sillas, sombrillas, living)", [
                "Vista general del sector de mesas y sillas (conjunto completo, orden y distribución).",
                "Detalle de una mesa: superficie de la tapa, ploteo, estado de pintura.",
                "Detalle de una silla: patas, regatones, estado de pintura, soldaduras.",
                "Sombrillas (si existen): foto general mostrando estado y estabilidad.",
                "Living (si existe): foto de sillones/tapizados mostrando estado del tapizado, lonas, estructura.",
            ], "Si hay defectos (óxido, pinturas descascaradas, tapizados rotos): foto extra en detalle."),
            ("A.4", "Iluminación interior: Estado y limpieza", [
                "Vista general del salón mostrando todas las luminarias (que se aprecie si están encendidas y si la luz es suficiente).",
                "Vista general del AMH mostrando luminarias.",
                "Detalle de un plafón/luminaria de cerca (para ver limpieza, presencia de insectos, estado).",
                "Foto de cualquier luminaria apagada o con defecto visible.",
            ], None),
            ("A.5", "Pisos, aberturas, techo y zócalos en salón: Estado y limpieza", [
                "Pisos del salón: vista general (que se vean las juntas, estado de las piezas, zócalos).",
                "Detalle de zócalos: acercar a la unión piso-pared.",
                "Aberturas/puertas de ingreso: foto frontal mostrando vidrios, herrajes, estado general.",
                "Detalle de vidrios: buscar rajaduras, falta de masilla, vinilos/fajas.",
                "Techo/cielorraso del salón: foto mirando hacia arriba.",
                "Alfombra de ingreso (si existe): foto de cerca mostrando estado.",
            ], "Si hay cortina metálica: foto cuando está abierta (buscar graffitis en cara interna)."),
            ("A.6", "Muros y revestimientos de muros en salón: Estado y limpieza", [
                "Pared principal del salón (revestimiento nebraska, cuadros, ploteos).",
                "Pared del AMH: cerámicos bricks, estado de juntas.",
                "Detalle de cuadros decorativos (bien puestos, imágenes autorizadas).",
                "Zócalos y repisas corian en unión de módulos con pared de bricks.",
                "Cualquier pared con humedad, pintura descascarada, cerámicos rotos o faltantes (foto de detalle).",
            ], None),
            ("A.7", "Cartelera: Estado y limpieza", [
                "Vista frontal completa de la cartelera (que se lean los precios y se vean todas las pantallas/láminas).",
                "Detalle de precios: foto acercada donde se verifique que son los autorizados.",
                "Sabores faltantes: foto mostrando si están marcados correctamente (sticker o marcación del sistema).",
                "Estado de limpieza de la cartelera: marcos, cables, estructura.",
            ], None),
            ("A.8", "Mostradores: Estado y limpieza", [
                "Vista frontal completa del mostrador (desde el lado del cliente).",
                "Detalle de la piedra del mostrador: buscar rajaduras, roturas.",
                "Vidrios del mostrador: estado, limpieza, fijación.",
                "Material de comunicación exhibido en el mostrador.",
                "Iluminación del mostrador (si aplica): foto mostrando si funciona.",
            ], None),
            ("A.9", "Separación de áreas / Orden y estado del depósito", [
                "Puerta de ingreso al depósito: señalización visible.",
                "Vista general del depósito (que se vea el orden, distribución, estanterías, tarimas).",
                "Paredes del depósito: cerámicos, pintura, estado de juntas.",
                "Pisos del depósito: foto de cerca mostrando uniones curvas pared-piso, estado general.",
                "Techo/cielorraso del depósito.",
                "Estanterías y tarimas: material, estado, separación de 5 cm respecto de la pared.",
                "Casilleros de colaboradores (interior si es posible).",
                "Tarima de recepción de mercadería: foto mostrando existencia y material.",
            ], "Si existe cámara de frío: foto adicional del interior mostrando piso sanitario, estantes, burletes."),
            ("A.10", "¿El local permite descarga eficiente de productos congelados?", [
                "Puerta de ingreso por donde pasa el carrito/zorrita (que se aprecie el ancho).",
                "Recorrido desde la puerta hasta el depósito (posibles obstáculos, escaleras, desniveles).",
                "Rampa (si existe) o escaleras con su señalización.",
                "Espacio de estacionamiento frente al local para el camión.",
            ], None),
            ("A.11", "¿Hay espacio disponible para descansar?", [
                "Vista general del espacio de descanso de colaboradores.",
                "Detalle del mobiliario (sillas/mesa) y limpieza del sector.",
                "Foto mostrando que no hay insumos ni mobiliario almacenado en ese espacio (o foto si los hay).",
            ], None),
        ]),
        ("B", "EXPERIENCIA DEL CLIENTE",
         "Cada 2 MESES (solo en auditoría COMPLETA)", [
            ("B.1", "Portal de Novedades, musicalización en exterior y en salón", [
                "TV del salón encendido mostrando el portal de novedades (que se vea la pantalla y el contenido).",
                "Equipo de música / parlantes: ubicación y estado.",
                "Cables de conexión: foto mostrando si están ordenados u ocultos.",
                "Vista general del salón donde se aprecie la ubicación del TV (verificar que NO esté en el AMH).",
            ], None),
            ("B.2", "Climatización y ventilación: Estado y limpieza", [
                "Equipo de aire acondicionado: vista general de la unidad interior.",
                "Detalle de rejillas y filtros del equipo (acercar para ver limpieza, polvo, suciedad).",
                "Manguera de desagote y bidón de recolección.",
                "Carcasa del equipo: buscar telas de araña, manchas de sarro.",
                "Si hay equipo exterior (compresor): foto mostrando estado.",
            ], "IMPORTANTE: si se usa calefactor, estufa o ventilador NO autorizado, fotografiar el equipo."),
            ("B.3", "Dispenser de agua", [
                "Dispenser completo: vista frontal mostrando ubicación en salón/AMH.",
                "Detalle de canillas y receptáculo/rejilla de goteo.",
                "Estado de higiene de la carcasa y bandeja inferior.",
                "Disponibilidad de vasos junto al dispenser.",
            ], None),
            ("B.4", "Matafuegos, luces de emergencia, instalación eléctrica y señalética", [
                "Matafuegos del salón: foto completa mostrando ubicación, accesibilidad y que está colgado.",
                "Detalle de la etiqueta de carga del matafuegos (fecha de vencimiento, tipo ABC, peso).",
                "Matafuegos del depósito (si hay otro).",
                "Luz de emergencia del salón: foto mostrando que está enchufada.",
                "Luz de emergencia del depósito.",
                "Tablero eléctrico: foto mostrando tapas sanas, señalética de riesgo eléctrico, llaves térmicas y disyuntor.",
                "Señalética de salida de emergencia.",
                "Desniveles/escalones: foto mostrando señalización en amarillo o rojo (o la falta de ella).",
            ], "Si hay barandas (escaleras, terrazas): foto adicional mostrando estado e instalación."),
            ("B.5", "Baño para clientes", [
                "Puerta del baño con señalización visible.",
                "Vista general del interior del baño.",
                "Inodoro: tapa completa, botón/mochila funcionando.",
                "Lavamanos: grifería, dispenser de jabón con contenido, dispenser de toallas con contenido.",
                "Espejo: estado, película antiestallido.",
                "Ventiluz/extractor: estado de la tela mosquitera.",
                "Secamanos eléctrico (si existe): enchufado y funcionando.",
            ], "Si hay cambiador para bebés: foto adicional mostrando estado."),
            ("B.6", "Sector de juegos infantiles", [
                "Vista general del sector de juegos (conjunto completo, piso blando, delimitación).",
                "Detalle de cada juego: estructura, fijaciones, pintura.",
                "Piso blando/goma: foto de cerca mostrando que cubre toda la superficie del juego.",
                "Estado de higiene: buscar pegotes de helado, suciedad acumulada.",
                "Maquinitas (si existen): foto mostrando estado y funcionamiento.",
            ], "NO APLICA si el local no tiene sector infantil."),
            ("B.7", "SmartFran - Club Grido - Conexión y funcionamiento", [
                "Pantalla de SmartFran / sistema loyalty mostrando que está logueado y funcionando.",
                "Lector de tarjetas Club Grido: foto del dispositivo.",
                "Tabla de canjes oficial impresa: foto mostrando que está exhibida y a la vista del cliente.",
            ], None),
            ("B.8", "APP GRIDO: Gestión del canal digital", [
                "Material de comunicación de descarga de la APP exhibido en el local (soporte fijo).",
                "Captura de pantalla de la APP mostrando que el local está disponible/activo en take away.",
                "Captura de pantalla del panel de pedidos mostrando pedidos efectivos y cancelados del último mes.",
            ], None),
            ("B.9", "Posnet y Grido Go (tótem autogestión)", [
                "Posnet: foto mostrando el equipo conectado y en funcionamiento.",
                "Tótem de autogestión (si existe): vista frontal completa mostrando pantalla encendida.",
                "Detalle del tótem: elementos completos, estado de higiene.",
            ], "NO APLICA el tótem si el local no lo tiene."),
            ("B.10", "Disponibilidad de Wi-Fi para el cliente", [
                "Señalética oficial de Wi-Fi exhibida en el local.",
                "Captura de pantalla del celular mostrando que la red Wi-Fi está disponible.",
                "Ubicación de la señalética en contexto (que se vea dónde está colocada).",
            ], None),
            ("B.11", "¿Dispone de generador eléctrico?", [
                "Generador eléctrico: vista general mostrando ubicación.",
                "Detalle del equipo: estado, conexiones.",
                "Si NO tiene generador: foto del tablero eléctrico general.",
            ], None),
        ]),
        ("C", "OPERATORIA DIARIA",
         "MENSUAL (en auditoría COMPLETA y RÁPIDA)", [
            ("C.1", "Disponibilidad de productos a granel", [
                "Vista general del pozo/mostrador de granel abierto mostrando todos los sabores disponibles.",
                "Foto panorámica de la cartelera de sabores (para cruzar contra lo que hay en el mostrador).",
                "Detalle de los huecos vacíos donde faltan sabores (si los hay).",
                "Foto de sabores marcados como faltantes en la cartelera.",
            ], None),
            ("C.2", "Disponibilidad y exhibición de productos secos e insumos generales", [
                "Zona de toppings: nueces, tramontana, frutilla, crispines, ensalada de frutas. Foto de conjunto.",
                "Salsas: foto mostrando variedad disponible (500g y 120g).",
                "Vasos, cucharitas, envases térmicos: foto del sector de insumos.",
                "Servilleteros en el salón: foto mostrando cantidad y distribución.",
                "Exhibidora de Grido Tops: foto frontal mostrando variedades, estado y vinilo/faja.",
                "Grido Market: foto frontal del exhibidor mostrando productos de frente y nivel de ocupación (3/4 mínimo).",
                "Otros exhibidores de productos temporales: foto mostrando precios autorizados.",
            ], None),
            ("C.3", "Condiciones de almacenamiento, fraccionamiento y rotulado", [
                "Estantería del depósito: vista general mostrando el orden de productos (limpieza abajo, materias primas medio, envases arriba).",
                "Productos sobre tarimas: foto que permita apreciar la altura respecto del piso (mínimo 14 cm).",
                "Detalle de etiquetas de productos cerrados: que se lean RNE, RNPA, fecha de elaboración/vencimiento.",
                "Sector de productos fallados/vencidos: foto de la cartelería que los identifica.",
                "Recipientes de productos fraccionados: foto mostrando cierre hermético, estado, etiqueta (producto, marca, fechas).",
                "Registro de productos fraccionados: foto del documento impreso o pantalla digital de los últimos 30 días.",
                "Interior de la cámara de frío (si aplica): foto mostrando estibado.",
            ], None),
            ("C.4", "Equipos de frío: Higiene y mantenimiento (pozos, cámaras, freezer, exhibidora, toppinera, frigobar)", [
                "Freezer del depósito: vista externa (ploteo, tapas, burletes).",
                "Freezer del depósito: interior abierto (hielo acumulado, limpieza).",
                "Pozo de frío del mostrador: vista superior abierta.",
                "Exhibidora vertical (impulsivos): foto frontal completa mostrando ploteo, cenefa, vidrios.",
                "Exhibidora horizontal (si existe): foto superior abierta.",
                "Toppinera: foto mostrando estado, tapas, limpieza.",
                "Frigobar: foto abierto mostrando interior (productos en tuppers cerrados y rotulados).",
                "Cámara de frío (si existe): burletes de puerta, cortina sanitaria, piso sanitario.",
            ], None),
            ("C.5", "Temperatura de equipos de frío", [
                "Termómetro del pozo de granel en mostrador mostrando temperatura (debe estar entre -13,5°C y -15°C).",
                "Termómetro del freezer del depósito mostrando temperatura (debe estar a -18°C).",
                "Termómetro de la cámara de frío (si aplica) mostrando -18°C.",
                "Termómetro de la exhibidora de impulsivos mostrando -18°C.",
                "Foto de helados a granel de cerca: verificar que no haya cristalización, deformación o pérdida de consistencia.",
            ], None),
            ("C.6", "Capacidad de frío en la franquicia", [
                "Vista general de todos los equipos de frío del depósito (cantidad total de freezers/cámaras).",
                "Detalle de la capacidad: foto del interior de un freezer representativo con nivel de ocupación.",
                "Cámara de frío (si existe): vista general del interior.",
            ], None),
            ("C.7", "Prevención de la contaminación cruzada", [
                "Cajas de helado a granel en el pozo: foto superior mostrando estado de las cajas (solapas, presentación, bajado).",
                "Detalle de 2-3 latas/cajas: buscar presencia de hielo, restos de cucuruchos u otros sabores.",
                "Interior de las cajas/baschetas: foto que muestre si están \"demasiado limpias\" (indicio de limpieza con esponja).",
                "Colaborador atendiendo: foto mostrando manipulación correcta (cucuruchos con servilleta, térmicos por la base).",
                "Lavamanos del AMH: foto mostrando que está accesible para lavado de manos frecuente.",
                "Tarima sanitaria (si se está reponiendo): foto mostrando cajas fuera del freezer sobre tarima.",
            ], None),
            ("C.8", "Abastecimiento de agua", [
                "Grifo del AMH: foto mostrando salida de agua fría y caliente.",
                "Grifo del depósito: ídem.",
                "Tanque de agua (si se tiene acceso): foto mostrando estado de limpieza.",
                "Foto del agua corriendo: verificar que no está turbia.",
            ], None),
            ("C.9", "AMH (mesadas, bacha, lavabochero, grifo monocomando)", [
                "Vista general del AMH completo (mesada, bacha, bajomesada).",
                "Detalle de la bacha: estado, limpieza, desagüe.",
                "Lavabochero: foto mostrando estado y limpieza del agua.",
                "Monocomando: foto de cerca mostrando estado y funcionamiento.",
                "Bajomesada abierto: foto mostrando interior (puertas, melamina, herrajes, ausencia de objetos personales).",
                "Mesada: estado de la piedra, limpieza, ausencia de elementos ajenos al sector.",
            ], None),
            ("C.10", "Equipos de producción (licuadora, cremera, chocolatera, maq. café, bols, utensilios)", [
                "Licuadora: foto completa mostrando vaso, base, cables.",
                "Chocolatera: foto mostrando nivel de chocolate, tapa, cuchara.",
                "Cremera: foto mostrando pico, estado de la rosca.",
                "Máquina de café: foto frontal, detalle de base y desagüe.",
                "Bols con tapa: foto mostrando estado y limpieza (incluida tapa).",
                "Utensilios alineados: foto de las 6 paletas, 6 bocheros y 4 corvetes juntos (para contar cantidad y verificar material).",
                "Detalle de mangos de utensilios: foto de cerca verificando que sean de acero inoxidable (no madera, no aluminio).",
                "Termómetro de punción: foto mostrando estado y limpieza.",
            ], None),
            ("C.11", "Equipos complementarios (balanza, microondas, caja, PC)", [
                "Balanza: foto mostrando visor a la vista del cliente, estado de la bandeja.",
                "Microondas: foto frontal mostrando display, puerta, plato.",
                "Microondas abierto: foto del interior mostrando limpieza.",
                "Sector de caja/PC: foto mostrando orden de cables y ausencia de papeles con vista al cliente.",
                "Vista desde el lado del cliente del sector caja: verificar que no hay papeles, notas o elementos no autorizados.",
            ], None),
            ("C.12", "Elementos de higiene personal (dispenser jabón y toallas en depósito y AMH)", [
                "Dispenser de jabón en AMH: foto mostrando estado, fijación a la pared, contenido.",
                "Dispenser de toallas en AMH: foto mostrando estado y contenido.",
                "Dispenser de jabón en depósito: foto ídem.",
                "Dispenser de toallas en depósito: foto ídem.",
            ], None),
            ("C.13", "Elementos de limpieza y desinfección. Ausencia de sustancias peligrosas", [
                "Sector de almacenamiento de productos de limpieza: foto mostrando ubicación, cartelería, separación de materia prima.",
                "Rejillas extendidas: foto mostrando las de cada color (blanco = AMH, amarillo = salón/exterior, otro = baño).",
                "Esponjas: foto mostrando tipo (verificar que NO sean metálicas en AMH).",
                "Escobillón, escoba, pala, secador, mopa: foto mostrando estado y tipo de cabo (plástico o metal forrado, NO madera).",
                "Productos químicos: foto mostrando envases cerrados, identificados con etiqueta original.",
                "Foto general del sector verificando que NO hay plaguicidas, venenos, nafta, solventes, insecticidas.",
                "Generador eléctrico en depósito (si aplica): verificar que tiene cobertura.",
            ], None),
            ("C.14", "¿El baño de los colaboradores está en condiciones?", [
                "Vista general del baño de colaboradores.",
                "Paredes y pisos: buscar hongos, humedad, suciedad.",
                "Techo: buscar telas de araña.",
                "Ventiluz/extractor con tela mosquitera.",
                "Insumos: papel higiénico, toalla de mano, jabón disponibles.",
            ], None),
            ("C.15", "Higiene y Salud del Colaborador", [
                "Foto de cada colaborador presente de frente (cuerpo completo): verificar uniforme, cofia, pelo recogido, ausencia de adornos.",
                "Detalle de manos de cada colaborador: uñas cortas, limpias, sin esmalte, sin anillos/pulseras.",
                "Colaborador manipulando producto: verificar buenas prácticas (cucurucho con servilleta, sin tocar con manos).",
                "Foto del lavamanos con jabón accesible en AMH.",
            ], "NOTA: pedir permiso al colaborador antes de fotografiarlo."),
            ("C.16", "Uniforme del Colaborador", [
                "Foto de cada colaborador de cuerpo completo: verificar cofia, delantal jean, remera, pantalón jean.",
                "Detalle del logo del uniforme (verificar que está actualizado).",
                "Calzado: foto de los pies mostrando calzado cerrado.",
                "Estado del uniforme: buscar roturas, remiendos, decoloración.",
            ], None),
            ("C.17", "Conocimiento del colaborador", [], "ESTE ÍTEM NO REQUIERE FOTOS. Se evalúa mediante pregunta verbal aleatoria al colaborador sobre productos o procedimientos. Anotar la pregunta realizada y la respuesta obtenida en un texto o audio."),
            ("C.18", "Constancia AFIP, Inscripción Provincial, Habilitación Municipal, Carnet Sanitario y Control de plagas", [
                "Constancia de Inscripción AFIP exhibida.",
                "Inscripción Provincial (Ingresos Brutos) exhibida.",
                "Habilitación Municipal exhibida.",
                "Carnet de manipulador de alimentos de cada colaborador presente (foto de cada carnet vigente).",
                "Certificado de desinfección exhibido (verificar que no esté vencido, con sello y firma) + formulario MIP.",
            ], None),
        ]),
        ("D", "IMAGEN: Formato y Estética",
         "Cada 2 MESES (solo en auditoría COMPLETA)", [
            ("D.1", "Exterior: Formato (pisos, zócalos, muros, iluminación, pérgolas, toldos)", [
                "Vista frontal completa del local desde la calle (colores de muros, techo, cajón de cortina metálica).",
                "Pisos exteriores: material y formato (verificar contra planilla técnica).",
                "Pérgola/toldo: foto mostrando formato y color.",
                "Motores de aire acondicionado en fachada: ubicación correcta.",
                "Color de muros exteriores y techos: foto que permita comparar con el color autorizado.",
            ], None),
            ("D.2", "Marquesina: formato", [
                "Marquesina completa de frente: verificar logo NUEVO.",
                "Saliente de pared: verificar imagen actualizada.",
                "Cartel bandera (si existe): verificar si está actualizado.",
                "Brazos reflectores: verificar que sean LED.",
            ], None),
            ("D.3", "Comunicación en vidriera", [
                "Vidriera completa de frente: que se vean todos los soportes de comunicación, vinilos y microperforados.",
                "Faja esmerilada: foto mostrando diseño y altura (90 cm del piso).",
                "Soportes de comunicación: foto de cada uno (máximo 2 permitidos).",
                "Puerta: foto mostrando carteles de servicio (abierto-cerrado, horarios, Club Grido, WiFi).",
                "Cartel \"se busca personal\" (si existe): verificar ubicación en puerta por sobre la faja y formato autorizado.",
                "Cartel Delivery (si existe): verificar diseño autorizado.",
            ], None),
            ("D.4", "Mobiliario exterior: formato (mesas, sillas, sombrillas, cerco)", [
                "Vista general de mesas y sillas exteriores: verificar modelo (ONE o Río) y colores.",
                "Detalle de una silla: verificar modelo (ONE: blanco/azul/rojo o naranja; o Río). Verificar que NO convivan rojas con naranjas.",
                "Sombrillas: color (blancas o azules), formato según ficha.",
                "Cerco o maceteros delimitando el sector (si aplica).",
                "Bancos (si existen): verificar modelo (metálicos azules o madera según ficha, que NO convivan).",
            ], "Si hay comunicación en mesas o servilleteros: foto de detalle."),
            ("D.5", "Mobiliario interior: formato (mesas, sillas, living)", [
                "Vista general del salón interior con mesas y sillas.",
                "Detalle del modelo de sillas interiores (ONE o Río; que NO convivan en el mismo sector).",
                "Barra con banquetas (si existe): verificar modelo MILO.",
                "Living (si existe): butacón BARI con mesa ratona de madera O sillón PAMPA con mesa ratona blanca.",
                "Verificar ausencia de bancos metálicos en el salón.",
            ], None),
            ("D.6", "Iluminación interior: formato", [
                "Luminarias del AMH: verificar tipo autorizado (plafones LED, riel blanco, galponeras, bandejas) y tono (frío o neutro).",
                "Luminarias del sitting/salón: verificar tipo y tono (cálido o neutro).",
                "Lámparas colgantes (si existen): verificar formato autorizado.",
                "Detalle de un artefacto de iluminación (modelo/marca).",
            ], None),
            ("D.7", "Pisos, aberturas, techo y zócalos en salón: Formato", [
                "Pisos del salón: verificar material y color según planilla técnica.",
                "Techo/cielorraso: verificar que es del mismo color que los muros.",
                "Iluminación de la garganta del AMH: verificar luces blancas, cálidas o neutras (NO de color).",
                "Cortinas (si existen): verificar que sean blackout, solo en sitting, de color autorizado.",
                "Alfombra de ingreso: estado y formato actualizado.",
            ], None),
            ("D.8", "Muros y revestimientos en salón: formato", [
                "Paredes del salón: verificar color autorizado.",
                "Ploteos: foto de cada ploteo verificando que estén actualizados.",
                "AMH: cerámicos bricks presentes y en condiciones.",
                "Revestimiento de madera/empapelado: verificar logo actualizado (no el viejo).",
                "Zócalos y repisas corian: presentes en la unión módulos-bricks.",
                "Caja de durlock cubriendo rollos de cortina metálica (si aplica): verificar que es del color de los muros.",
            ], None),
            ("D.9", "Cartelera: formato", [
                "Cartelera completa: verificar formato digital vigente.",
                "TVs: verificar tamaño (43 pulgadas), alineación entre pantallas.",
                "Foto desde la caja: verificar legibilidad de la primera lámina.",
                "Sabores faltantes marcados en el sistema (si aplica).",
            ], None),
            ("D.10", "Mostradores: formato", [
                "Mostrador completo de frente: verificar modelo nebraska/báltico.",
                "Detalle del logo en el mostrador (verificar actualizado).",
                "Material de comunicación en el mostrador: verificar autorizado.",
                "Vidrios e iluminación del mostrador.",
            ], None),
            ("D.11", "Decoración, plantas, jardines y cestos", [
                "Plantas del salón y exterior: estado y distribución.",
                "Decoraciones de fecha especial (si aplica): terminación y completitud.",
                "Cestos del salón: modelo autorizado (plástico tapa vaivén).",
                "Cestos exteriores: modelo autorizado (chapa).",
                "Cesto del baño: verificar modelo a pedal.",
            ], "Si hay soporte de vereda/cartel sandwich: foto verificando que NO esté escrito a mano."),
            ("D.12", "Sector de juegos infantiles (formato)", [
                "Vista general del sector: empapelado que define el sector, delimitación.",
                "Juegos: verificar estado de seguridad y formato.",
                "Piso blando interior o césped exterior.",
                "Maquinitas (máximo 1): verificar ubicación en sector infantil.",
            ], "NO APLICA si el local no tiene sector infantil."),
            ("D.13", "Circuito ingreso-compra-consumo-salida: ¿es claro y cómodo?", [
                "Desde la puerta de entrada mirando hacia adentro: que se vea el AMH enfrentado a la puerta.",
                "Recorrido del cliente: foto mostrando la secuencia natural desde el ingreso hacia el mostrador.",
                "Sector de autoservicio: verificar que está antes o junto a caja.",
                "Sector infantil: verificar que está fuera de zona de riesgo.",
            ], None),
            ("D.14", "¿La cantidad de puestos de trabajo y cajas es coherente?", [
                "Vista del AMH completo mostrando la cantidad de puestos de trabajo operativos.",
                "Sector de caja(s): foto mostrando cuántas cajas hay.",
                "Ventana take away (si existe para delivery de alta demanda).",
            ], None),
            ("D.15", "¿La franquicia se encuentra ordenada?", [
                "Vista panorámica del salón: distanciamiento entre mesas.",
                "Identificación de sectores: infantil, living, sitting general.",
                "Servicios complementarios implementados.",
                "Vista general desde la puerta mostrando el orden general.",
            ], None),
            ("D.16", "En términos generales, ¿el local cumple con los requisitos de marca?", [
                "Fachada completa del local desde la vereda de enfrente.",
                "Vista panorámica interior del salón completo.",
                "Vista panorámica del AMH completo.",
            ], None),
        ]),
        ("E", "OFERTA Y STOCK",
         "Cada 2 MESES (solo en auditoría COMPLETA)", [
            ("E.1", "Disponibilidad y exhibición de impulsivos", [
                "Exhibidora vertical: foto frontal completa (verificar ploteo vinilo simil acero inoxidable, productos de cara al cliente, ocupación mínima 3/4).",
                "Exhibidora horizontal (si existe): foto superior abierta (productos de cara, cajas cerradas, canastos con sueltas).",
                "Detalle del ploteo/vinilo: estado, marca (Mc Cal 6076).",
                "Cenefa de la exhibidora: estado y precios autorizados.",
                "Vinilos de promociones o señalética: estado y autorización.",
                "Detalle de productos faltantes o espacios vacíos.",
            ], "Verificar que NO hay pozos o freezers operativos en el salón."),
            ("E.2", "Disponibilidad y exhibición de congelados (Frizzio)", [
                "Exhibidora exclusiva Frizzio: foto frontal mostrando todas las familias de productos exhibidas.",
                "Cenefa Frizzio (horizontales) o vinilo (verticales): estado y actualización.",
                "Pieza de comunicación Frizzio en vidriera (A1) o sector caja (A5).",
                "Precios en vinilo (verificar que NO son en papel ni a mano).",
                "Nivel de ocupación de la exhibidora (mínimo 3/4).",
            ], None),
            ("E.3", "¿Posee cámara de frío?", [
                "Cámara de frío: vista exterior mostrando puerta y ubicación.",
                "Interior de la cámara: vista general.",
                "Si está fuera de la franquicia: foto mostrando ubicación/acceso.",
            ], "Si NO posee cámara de frío: 1 foto del depósito mostrando que no existe."),
            ("E.4", "¿La cámara de frío tiene estantes?", [
                "Interior de la cámara: foto mostrando estantes (si existen).",
                "Palets/piso sanitario: foto de detalle.",
                "Combinación de estantes y palets (si aplica).",
            ], "NO APLICA si el local no tiene cámara de frío."),
            ("E.5", "¿Cumple con capacidad de frío ideal?", [
                "Todos los equipos de frío alineados o en vista general (para estimar capacidad total).",
                "Etiqueta/placa de cada freezer o equipo mostrando capacidad en litros.",
                "Interior de equipos mostrando nivel de ocupación actual.",
            ], None),
        ]),
    ]

    for sec_letter, sec_name, periodicity, items in sections_data:
        h_sec = doc.add_heading(f"Sección {sec_letter} — {sec_name}", level=1)
        for run in h_sec.runs:
            run.font.color.rgb = RED

        p_per = doc.add_paragraph()
        rp = p_per.add_run(f"PERIODICIDAD: {periodicity}")
        rp.bold = True
        rp.font.size = Pt(10)
        rp.font.color.rgb = DARK_BLUE
        p_per.space_after = Pt(4)

        for item_id, item_name, photos, tip in items:
            add_item_block(doc, item_id, item_name, photos, tip)

        doc.add_page_break()

    # ── Tips para el colaborador ──
    h_tips = doc.add_heading("Tips para el colaborador", level=1)
    for run in h_tips.runs:
        run.font.color.rgb = RED

    tips = [
        "Sacá las fotos con buena iluminación y sin flash directo.",
        "No uses filtros ni editores de imagen.",
        "Intentá que las fotos estén bien enfocadas y sin movimiento.",
        "Nombrá cada foto con el código del ítem (A1_01, A1_02, etc.) o agrupalas en carpetas por sección (A, B, C, D, E).",
        "Si encontrás un defecto: sacale una foto general y una de detalle.",
        "Las fotos de documentos (carnets, certificados, registros) deben ser legibles: acercar lo suficiente para que se lean las fechas.",
        "Las fotos de termómetros deben mostrar la lectura con claridad.",
        "Si un ítem \"No Aplica\" a tu local, anotalo y no saques fotos.",
        "Aprovechá el recorrido natural: empezá por afuera (A.1, A.2), entrá al salón (A.3-A.8), seguí al depósito (A.9), y luego recorré los demás ítems en orden.",
    ]
    for t in tips:
        p = doc.add_paragraph(t, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(9)

    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build()
